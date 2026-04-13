#!/usr/bin/env python3
"""
Phase 1 migration: 既存 TAGS (compiled/tags.json) + SAFE 未収録 103 items を
v2 JSON スキーマ (categories > sections > tags) に変換する。

出力: data/v2/tags.json
"""

import json
import datetime
from pathlib import Path
import re

ROOT = Path(__file__).parent.parent

# ---------------------------------------------------------------------------
# 1. セクション → カテゴリ マッピング
# ---------------------------------------------------------------------------

# 既存 TAGS の 79 セクション(label) を v2 category id へマップ
SECTION_TO_CATEGORY = {
    # カメラ・構図
    "camera_comp":   "camera",
    "focus":         "camera",
    "angle":         "camera",
    "gaze":          "camera",
    "camera":        "camera",
    "pov":           "camera",

    # 服装
    "clothes":       "clothing",

    # 人数・関係
    "count":         "people",
    "relationship":  "people",
    "misc_people":   "people",
    "layout":        "people",

    # メタ・品質
    "tech2":         "meta",
    "quality":       "meta",
    "effect":        "meta",
    "style":         "meta",
    "manga_panel":   "meta",
    "manga_read":    "meta",
    "cover":         "meta",
    "frame":         "meta",
    "meta_text":     "meta",
    "qc":            "meta",
    "artifact":      "meta",
    "revision":      "meta",
    "dup":           "meta",
    "shape":         "meta",
    "orientation":   "meta",
    "quality2":      "meta",
    "effect2":       "meta",
    "style2":        "meta",
    "doc":           "meta",
    "ui":            "meta",

    # 表情・顔
    "expr_smile":    "expression",
    "expr_evil":     "expression",
    "eye":           "expression",
    "eye_empty":     "expression",
    "mouth":         "expression",
    "expr_misc":     "expression",
    "shake":         "expression",
    "tired":         "expression",
    "mood_bad":      "expression",
    "anger":         "expression",
    "confusion":     "expression",
    "anxiety":       "expression",
    "fear":          "expression",
    "sad":           "expression",
    "teeth":         "expression",
    "nose":          "expression",
    "brow":          "expression",
    "face_feature":  "expression",
    "makeup":        "expression",
    "lip":           "expression",
    "paint":         "expression",
    "mark":          "expression",

    # ポーズ
    "pose_hand":     "pose",
    "pose_arm":      "pose",
    "touch_self":    "pose",
    "touch_env":     "pose",
    "rest":          "pose",
    "point":         "pose",
    "support":       "pose",
    "misc_pose":     "pose",
    "sit":           "pose",
    "sexpos":        "pose",
    "legs":          "pose",
    "feet":          "pose",
    "kneel":         "pose",
    "stand":         "pose",
    "torso":         "pose",
    "legmove":       "pose",
    "other_pose":    "pose",
    "uncat":         "pose",

    # 動作・行動
    "action":        "action",
    "trouble":       "action",
    "hold":          "action",
    "prep":          "action",
    "body":          "action",
    "pose_action":   "action",

    # アクセサリー
    "accessories":   "accessories",

    # E621・Pony
    "e621_pony":     "e621_pony",
}

# セクションの日本語ラベル
SECTION_LABELS = {
    "camera_comp":  "構図",
    "focus":        "フォーカス",
    "angle":        "アングル",
    "gaze":         "視線",
    "camera":       "カメラ技法",
    "pov":          "主観・POV",
    "clothes":      "服装・衣類",
    "count":        "人数",
    "relationship": "関係性",
    "misc_people":  "その他の人物",
    "layout":       "レイアウト",
    "tech2":        "テクニック",
    "quality":      "クオリティ",
    "effect":       "エフェクト",
    "style":        "スタイル",
    "manga_panel":  "マンガコマ",
    "manga_read":   "読み方向",
    "cover":        "カバー",
    "frame":        "フレーム",
    "meta_text":    "テキスト",
    "qc":           "QC",
    "artifact":     "アーティファクト",
    "revision":     "修正",
    "dup":          "複製",
    "shape":        "形状",
    "orientation":  "向き",
    "quality2":     "クオリティ2",
    "effect2":      "エフェクト2",
    "style2":       "スタイル2",
    "doc":          "ドキュメント",
    "ui":           "UI",
    "expr_smile":   "笑顔",
    "expr_evil":    "邪悪な笑み",
    "eye":          "目",
    "eye_empty":    "虚ろな目",
    "mouth":        "口",
    "expr_misc":    "その他の表情",
    "shake":        "震え・動揺",
    "tired":        "疲れ",
    "mood_bad":     "不機嫌",
    "anger":        "怒り",
    "confusion":    "混乱",
    "anxiety":      "不安",
    "fear":         "恐怖",
    "sad":          "悲しみ",
    "teeth":        "歯",
    "nose":         "鼻",
    "brow":         "眉",
    "face_feature": "顔の特徴",
    "makeup":       "メイク",
    "lip":          "唇",
    "paint":        "塗り",
    "mark":         "マーク",
    "pose_hand":    "手のポーズ",
    "pose_arm":     "腕のポーズ",
    "touch_self":   "自己接触",
    "touch_env":    "環境への接触",
    "rest":         "休憩・支え",
    "point":        "指し示す",
    "support":      "支持・サポート",
    "misc_pose":    "その他のポーズ",
    "sit":          "座る",
    "sexpos":       "性的ポーズ",
    "legs":         "脚",
    "feet":         "足",
    "kneel":        "跪く",
    "stand":        "立つ",
    "torso":        "上体",
    "legmove":      "脚の動き",
    "other_pose":   "その他のポーズ2",
    "uncat":        "未分類",
    "action":       "アクション",
    "trouble":      "トラブル・困り",
    "hold":         "保持・持つ",
    "prep":         "準備・姿勢",
    "body":         "体の動作",
    "pose_action":  "ポーズ＋アクション",
    "accessories":  "アクセサリー",
    "e621_pony":    "E621・Pony",
}

# カテゴリ定義（id / label / 順序）
CATEGORIES_DEF = [
    {"id": "camera",      "label": "カメラ・構図"},
    {"id": "expression",  "label": "表情・顔"},
    {"id": "pose",        "label": "ポーズ"},
    {"id": "action",      "label": "動作・行動"},
    {"id": "clothing",    "label": "服装"},
    {"id": "accessories", "label": "アクセサリー"},
    {"id": "people",      "label": "人数・関係"},
    {"id": "meta",        "label": "メタ・品質"},
    {"id": "e621_pony",   "label": "E621・Pony"},
]

# ---------------------------------------------------------------------------
# 2. target 自動判定（pose / action カテゴリのみ）
# ---------------------------------------------------------------------------

OTHER_KEYWORDS = [
    r"\banother'?s\b", r"\bsomeone\b", r"\bpartner\b",
    r"\bpatting\b", r"\bshoulder of\b", r"\bhug from behind\b",
]
MUTUAL_KEYWORDS = [
    r"\bholding hands\b", r"\barm.?link\b", r"\bsymmetrical docking\b",
    r"\binterlocked\b", r"\bface.?to.?face\b",
]
OBJECT_KEYWORDS = [
    r"\bholding\b", r"\bcarrying\b", r"\bwearing\b", r"\bwielding\b",
    r"\bgrabbing\b", r"\bphone\b", r"\bweapon\b", r"\bbottle\b",
    r"\bcup\b", r"\bbook\b", r"\bumbrella\b", r"\bsword\b",
    r"\bgun\b", r"\bwand\b", r"\bstaff\b",
]
# "self" に倒すキーワード（self hug, arms behind head など）
SELF_KEYWORDS = [
    r"\bself\b", r"\bown\b", r"\bbehind head\b", r"\bbehind back\b",
    r"\barms crossed\b", r"\bhand on hip\b", r"\bhands on\b",
    r"\btouch own\b",
]

TARGET_CATEGORIES = {"pose", "action"}


def infer_target(en: str, category_id: str):
    """en タグ文字列から target を推定。対象外カテゴリは (None, None) を返す。"""
    if category_id not in TARGET_CATEGORIES:
        return None, None

    en_l = en.lower()

    for pat in MUTUAL_KEYWORDS:
        if re.search(pat, en_l):
            return "mutual", None

    for pat in OTHER_KEYWORDS:
        if re.search(pat, en_l):
            return "other", None

    for pat in OBJECT_KEYWORDS:
        if re.search(pat, en_l):
            return "object", None

    # pose カテゴリは残りを self と見なす
    if category_id == "pose":
        return "self", None

    # action カテゴリは判定不能なら null
    for pat in SELF_KEYWORDS:
        if re.search(pat, en_l):
            return "self", None

    return None, None


# ---------------------------------------------------------------------------
# 3. SAFE 未収録 103 items のセクション割り当て
# ---------------------------------------------------------------------------

# SAFE category → v2 category_id / section_id のマップ
SAFE_CAT_TO_SECTION = {
    "action":     ("action",     "safe_action",     "SAFE: アクション"),
    "count":      ("people",     "count",           None),   # TAGS の count セクションへ追記
    "expression": ("expression", "safe_expression",  "SAFE: 表情（未分類）"),
    "focus":      ("camera",     "focus",           None),   # TAGS の focus セクションへ追記
    "pose":       ("pose",       "safe_pose",       "SAFE: ポーズ（未分類）"),
    "pov":        ("camera",     "pov",             None),   # TAGS の pov セクションへ追記
}

# ---------------------------------------------------------------------------
# 4. メイン処理
# ---------------------------------------------------------------------------

def last_japanese_pos(s: str) -> int:
    """文字列中の最後の日本語文字の位置を返す。なければ -1。"""
    for i in range(len(s) - 1, -1, -1):
        cp = ord(s[i])
        if (0x3000 <= cp <= 0x9fff) or (0xff00 <= cp <= 0xffef):
            return i
    return -1


def load_data2_docx() -> list:
    """data2/プロンプトCSVデータベース.docx から {en, jp, cat, desc} リストを返す。"""
    try:
        import docx as _docx
    except ImportError:
        print("  [WARN] python-docx not installed. data2 skipped.")
        return []

    path = ROOT / "data2/プロンプトCSVデータベース.docx"
    if not path.exists():
        print(f"  [WARN] {path} not found. data2 skipped.")
        return []

    doc = _docx.Document(str(path))
    raw = doc.paragraphs[0].text if doc.paragraphs else ""

    CATS = [
        "カメラ・構図", "ポーズ・動作", "接触動作（他者）", "接触動作（自分）",
        "接触動作（限定なし）", "服装・露出", "身体特徴", "アクセサリー・武器",
    ]
    cat_pattern = "|".join(re.escape(c) for c in CATS)
    cat_matches = list(re.finditer(r",(" + cat_pattern + r"),", raw))

    records = []
    for i, m in enumerate(cat_matches):
        cat = m.group(1)
        cat_start = m.start()
        cat_end = m.end()

        before = raw[:cat_start]
        last_comma = before.rfind(",")
        if last_comma == -1:
            continue
        jp = before[last_comma + 1 :].strip()

        if i > 0:
            prev_end = cat_matches[i - 1].end()
            between = raw[prev_end:last_comma]
            jp_last = last_japanese_pos(between)
            if jp_last >= 0:
                rest = between[jp_last + 1 :]
                space = rest.find(" ")
                en = rest[space + 1 :].strip() if space >= 0 else rest.strip()
            else:
                space_pos = between.rfind(" ")
                en = between[space_pos + 1 :].strip() if space_pos >= 0 else between.strip()
        else:
            between = before[:last_comma]
            jp_last = last_japanese_pos(between)
            if jp_last >= 0:
                rest = between[jp_last + 1 :]
                space = rest.find(" ")
                en = rest[space + 1 :].strip() if space >= 0 else rest.strip()
            else:
                space_pos = between.rfind(" ")
                en = between[space_pos + 1 :].strip() if space_pos >= 0 else between.strip()

        # desc
        if i < len(cat_matches) - 1:
            next_start = cat_matches[i + 1].start()
            segment = raw[cat_end:next_start]
            last_comma_d = segment.rfind(",")
            desc = segment[:last_comma_d].strip() if last_comma_d >= 0 else segment.strip()
        else:
            desc = raw[cat_end:].strip()

        if en:
            records.append({"en": en, "jp": jp, "cat": cat, "desc": desc})

    return records


def load_tags() -> dict:
    path = ROOT / "data/dictionary/compiled/tags.json"
    with open(path) as f:
        return json.load(f)

def load_safe() -> dict:
    path = ROOT / "data/dictionary/compiled/safe.json"
    with open(path) as f:
        return json.load(f)


def build_v2() -> dict:
    tags_data = load_tags()
    safe_data = load_safe()

    # ---- TAGS アイテムを section ごとに収集 ----
    section_items: dict[str, list] = {}
    for cat in tags_data["categories"]:
        key = cat["label"]
        raw_items = tags_data["items"][cat["start"]:cat["end"]]
        section_items[key] = raw_items

    # ---- カテゴリ構造を構築 ----
    cat_sections: dict[str, list] = {c["id"]: [] for c in CATEGORIES_DEF}

    for sec_key, raw_items in section_items.items():
        cat_id = SECTION_TO_CATEGORY.get(sec_key)
        if cat_id is None:
            print(f"  [WARN] unmapped section: {sec_key}")
            continue

        tags_list = []
        for item in raw_items:
            en = item.get("en", "")
            jp = item.get("jp", "")
            target, target_note = infer_target(en, cat_id)
            tags_list.append({
                "en": en,
                "jp": jp,
                "target": target,
                "target_note": target_note,
            })

        cat_sections[cat_id].append({
            "id": sec_key,
            "label": SECTION_LABELS.get(sec_key, sec_key),
            "tags": tags_list,
        })

    # ---- SAFE 未収録アイテムを TAGS en セットと照合して抽出 ----
    tags_en_set = {item["en"].lower().strip() for item in tags_data["items"]}

    # TAGS セクション id → セクション dict の参照マップ（追記用）
    sec_ref: dict[str, dict] = {}
    for cat_id, sections in cat_sections.items():
        for sec in sections:
            sec_ref[sec["id"]] = sec

    # 新規セクションのバッファ（safe_action, safe_expression, safe_pose）
    new_sections: dict[str, dict] = {}

    for safe_cat in safe_data["categories"]:
        safe_cat_key = safe_cat["key"]
        mapping = SAFE_CAT_TO_SECTION.get(safe_cat_key)
        if mapping is None:
            continue  # body_features / accessories / clothing / meta etc. は既に TAGS に存在

        cat_id, sec_id, new_sec_label = mapping

        for item in safe_cat["items"]:
            en = item.get("en", "")
            if en.lower().strip() in tags_en_set:
                continue  # 既に TAGS に存在

            # desc を jp に補足として付与
            jp = item.get("jp", "")
            desc = item.get("desc", "")
            if desc and desc not in jp:
                jp_combined = f"{jp}（{desc}）" if jp else desc
            else:
                jp_combined = jp

            target, target_note = infer_target(en, cat_id)
            tag_obj = {
                "en": en,
                "jp": jp_combined,
                "target": target,
                "target_note": target_note,
            }

            if new_sec_label is None:
                # 既存セクションへ追記
                if sec_id in sec_ref:
                    sec_ref[sec_id]["tags"].append(tag_obj)
                    tags_en_set.add(en.lower().strip())
            else:
                # 新規セクションへ格納
                if sec_id not in new_sections:
                    new_sections[sec_id] = {
                        "id": sec_id,
                        "label": new_sec_label,
                        "tags": [],
                        "_cat_id": cat_id,
                    }
                new_sections[sec_id]["tags"].append(tag_obj)
                tags_en_set.add(en.lower().strip())

    # 新規セクションをカテゴリへ追加
    for sec in new_sections.values():
        cat_id = sec.pop("_cat_id")
        cat_sections[cat_id].append(sec)

    # ---- data2 docx アイテムを追加 ----
    tags_en_set = {item["en"].lower().strip() for item in tags_data["items"]}
    # v2 に既に追加されたものも含めて再構築
    for c_id, sections in cat_sections.items():
        for sec in sections:
            for t in sec["tags"]:
                tags_en_set.add(t["en"].lower().strip())

    data2_records = load_data2_docx()

    # data2 カテゴリ → (v2_cat_id, sec_id, sec_label, fixed_target)
    DATA2_CAT_MAP = {
        "カメラ・構図":       ("camera",      "data2_camera",       "data2: カメラ・構図",     None),
        "ポーズ・動作":       ("pose",        "data2_pose",         "ポーズ・動作",            "self"),
        "接触動作（他者）":   ("action",      "data2_touch_other",  "接触動作（他者）",        "other"),
        "接触動作（自分）":   ("action",      "data2_touch_self",   "接触動作（自分）",        "self"),
        "接触動作（限定なし）": ("action",    "data2_touch_any",    "接触動作（限定なし）",    None),
        "服装・露出":         ("clothing",    "data2_clothing",     "服装・露出",              None),
        "身体特徴":           ("people",      "data2_body_features","身体特徴",                None),
        "アクセサリー・武器": ("accessories", "data2_accessories",  "アクセサリー・武器",      "object"),
    }

    data2_new_sections: dict[str, dict] = {}

    for rec in data2_records:
        en = rec["en"]
        if en.lower().strip() in tags_en_set:
            continue  # 既に v2 に存在

        d2_cat = rec["cat"]
        mapping = DATA2_CAT_MAP.get(d2_cat)
        if mapping is None:
            print(f"  [WARN] unmapped data2 cat: {d2_cat}")
            continue

        v2_cat_id, sec_id, sec_label, fixed_target = mapping

        if fixed_target is not None:
            target, target_note = fixed_target, None
        else:
            target, target_note = infer_target(en, v2_cat_id)

        jp = rec.get("jp", "")
        desc = rec.get("desc", "")
        if desc and desc not in jp:
            jp = f"{jp}（{desc}）" if jp else desc

        tag_obj = {"en": en, "jp": jp, "target": target, "target_note": target_note}

        if sec_id not in data2_new_sections:
            data2_new_sections[sec_id] = {
                "id": sec_id,
                "label": sec_label,
                "tags": [],
                "_cat_id": v2_cat_id,
            }
        data2_new_sections[sec_id]["tags"].append(tag_obj)
        tags_en_set.add(en.lower().strip())

    for sec in data2_new_sections.values():
        cat_id = sec.pop("_cat_id")
        cat_sections[cat_id].append(sec)

    # ---- v2 JSON 組み立て ----
    now = datetime.datetime.utcnow().isoformat(timespec="seconds") + "Z"
    categories = []
    for cat_def in CATEGORIES_DEF:
        cat_id = cat_def["id"]
        sections = cat_sections[cat_id]
        total = sum(len(s["tags"]) for s in sections)
        categories.append({
            "id": cat_id,
            "label": cat_def["label"],
            "sections": sections,
            "_count": total,
        })

    total_tags = sum(c["_count"] for c in categories)
    for c in categories:
        del c["_count"]

    return {
        "schema_version": "2.0",
        "generated_at": now,
        "count": total_tags,
        "categories": categories,
    }


def main():
    out_dir = ROOT / "data/v2"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "tags.json"

    print("Building v2 JSON...")
    v2 = build_v2()

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(v2, f, ensure_ascii=False, indent=2)

    print(f"Written: {out_path}")
    print(f"  schema_version : {v2['schema_version']}")
    print(f"  total tags     : {v2['count']}")
    print(f"  categories     : {len(v2['categories'])}")
    for cat in v2["categories"]:
        n_sections = len(cat["sections"])
        n_tags = sum(len(s["tags"]) for s in cat["sections"])
        print(f"    {cat['id']:15s}  {n_sections:3d} sections  {n_tags:5d} tags")


if __name__ == "__main__":
    main()
